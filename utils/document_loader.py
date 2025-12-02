"""
ContractGuard AI - 문서 로더 모듈
PDF, DOCX, TXT 파일에서 텍스트 추출
"""
import io
from typing import Optional

try:
    from PyPDF2 import PdfReader
except ImportError as e:
    raise ImportError(
        "PyPDF2가 설치되지 않았습니다. 다음 명령어로 설치하세요:\n"
        "pip install PyPDF2\n"
        f"원본 오류: {e}"
    )

try:
    from docx import Document
except ImportError as e:
    raise ImportError(
        "python-docx가 설치되지 않았습니다. 다음 명령어로 설치하세요:\n"
        "pip install python-docx\n"
        f"원본 오류: {e}"
    )


class DocumentLoader:
    """계약서 문서 로더"""
    
    @staticmethod
    def load_pdf(file) -> str:
        """PDF 파일에서 텍스트 추출"""
        try:
            if hasattr(file, 'read'):
                pdf_reader = PdfReader(io.BytesIO(file.read()))
            else:
                pdf_reader = PdfReader(file)
            
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"PDF 파일 읽기 오류: {str(e)}")
    
    @staticmethod
    def load_docx(file) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            if hasattr(file, 'read'):
                doc = Document(io.BytesIO(file.read()))
            else:
                doc = Document(file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"DOCX 파일 읽기 오류: {str(e)}")
    
    @staticmethod
    def load_txt(file) -> str:
        """TXT 파일에서 텍스트 추출"""
        try:
            if hasattr(file, 'read'):
                content = file.read()
                if isinstance(content, bytes):
                    return content.decode('utf-8').strip()
                return content.strip()
            else:
                with open(file, 'r', encoding='utf-8') as f:
                    return f.read().strip()
        except Exception as e:
            raise ValueError(f"TXT 파일 읽기 오류: {str(e)}")
    
    @classmethod
    def load(cls, file, file_type: Optional[str] = None) -> str:
        """파일 타입에 따라 자동으로 적절한 로더 선택"""
        if file_type is None:
            if hasattr(file, 'name'):
                file_type = file.name.split('.')[-1].lower()
            else:
                raise ValueError("파일 타입을 지정해주세요.")
        
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return cls.load_pdf(file)
        elif file_type in ['docx', 'doc']:
            return cls.load_docx(file)
        elif file_type == 'txt':
            return cls.load_txt(file)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {file_type}")

